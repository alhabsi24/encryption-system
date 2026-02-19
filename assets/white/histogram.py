import numpy as np
import cv2
import hashlib
import os
import pandas as pd
import matplotlib.pyplot as plt
from scipy.stats import entropy as scipy_entropy
from math import log10, sqrt
from docx import Document
from docx.shared import Inches

# ---------- DNA Encoding/Decoding ----------
dna_rules = {'00': 'A', '01': 'C', '10': 'G', '11': 'T'}
dna_decode = {v: k for k, v in dna_rules.items()}

def save_results_to_word(df, filename="results.docx"):
    doc = Document()
    doc.add_heading('Image Encryption Analysis Results', 0)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for idx, row in df.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    doc.save(filename)
    print(f"Table saved as {filename}")

def binary_to_dna(bin_str):
    return ''.join(dna_rules[bin_str[i:i+2]] for i in range(0, len(bin_str), 2))

def dna_to_binary(dna_str):
    return ''.join(dna_decode[n] for n in dna_str)

def dna_xor(base1, base2):
    b1 = dna_decode[base1]
    b2 = dna_decode[base2]
    xor_val = int(b1, 2) ^ int(b2, 2)
    return dna_rules[f'{xor_val:02b}']

def generate_lorenz_chaos(length, x0=0.1, y0=0, z0=0):
    s, r, b = 10, 28, 2.667
    dt = 0.01
    xs = []
    x, y, z = x0, y0, z0
    for _ in range(length):
        dx = s * (y - x)
        dy = r * x - y - x * z
        dz = x * y - b * z
        x += dx * dt
        y += dy * dt
        z += dz * dt
        xs.append(abs(x % 1))
    return np.array(xs)

def generate_salt(seed_length=16):
    salt = os.urandom(seed_length)
    hashed_salt = hashlib.sha256(salt).digest()
    seed_values = [int.from_bytes(hashed_salt[i:i+4], 'big') / (2**32) for i in range(0, 12, 4)]
    return salt, seed_values

def encrypt_image_with_dna(img_path):
    img = cv2.imread(img_path, cv2.IMREAD_UNCHANGED)
    if img is None:
        raise FileNotFoundError(f"Image not found at {img_path}")

    # Handle grayscale, binary, and color images
    if len(img.shape) == 2:  # Grayscale or binary
        h, w = img.shape
        ch = 1
        img_to_flat = img.flatten()
    elif len(img.shape) == 3:
        h, w, ch = img.shape
        img_to_flat = img.flatten()
    else:
        raise ValueError('Unknown image shape.')

    salt, (x0, y0, z0) = generate_salt()
    chaos_for_dna = generate_lorenz_chaos(len(img_to_flat) * 4, x0=x0, y0=y0, z0=z0)
    chaos_bin = ''.join([f'{int(c*256):08b}' for c in chaos_for_dna])
    chaos_bin = chaos_bin[:len(img_to_flat) * 8]
    dna_key = binary_to_dna(chaos_bin)
    img_bin = ''.join([f'{p:08b}' for p in img_to_flat])
    img_dna = binary_to_dna(img_bin)
    encrypted_dna = ''.join([dna_xor(img_dna[i], dna_key[i]) for i in range(len(img_dna))])
    encrypted_bin = dna_to_binary(encrypted_dna)
    encrypted_bytes = np.array([int(encrypted_bin[i:i+8], 2) for i in range(0, len(encrypted_bin), 8)], dtype=np.uint8)
    if ch == 1:
        encrypted_img = encrypted_bytes.reshape(h, w)
    else:
        encrypted_img = encrypted_bytes.reshape(h, w, ch)
    
    # Save encrypted image
    cv2.imwrite('encrypted_image.png', encrypted_img)
    
    return encrypted_img, img

# ---------- Metrics ----------
def calculate_entropy(image):
    if len(image.shape) == 2:
        histogram, _ = np.histogram(image.flatten(), bins=256, range=(0, 255), density=True)
        histogram = histogram + 1e-12
        return scipy_entropy(histogram, base=2)
    else:
        values = []
        for c in range(image.shape[2]):
            histogram, _ = np.histogram(image[:,:,c].flatten(), bins=256, range=(0, 255), density=True)
            histogram = histogram + 1e-12
            values.append(scipy_entropy(histogram, base=2))
        return np.mean(values)

def PSNR(img1, img2):
    mse = np.mean((img1.astype(np.float64) - img2.astype(np.float64)) ** 2)
    if mse == 0:
        return float('inf')
    max_pixel = 255.0
    return 20 * log10(max_pixel / sqrt(mse))

def NPCR(img1, img2):
    return np.sum(img1 != img2) / img1.size * 100

def UACI(img1, img2):
    return np.mean(np.abs(img1.astype(np.int16) - img2.astype(np.int16)) / 255) * 100

def correlation_coeff(img, direction='horizontal', n_points=3000):
    if len(img.shape) == 2:
        h, w = img.shape
        chs = 1
    else:
        h, w, chs = img.shape
    np.random.seed(0)
    result = []
    for channel in range(chs):
        values1, values2 = [], []
        for _ in range(n_points):
            if direction == 'horizontal':
                row = np.random.randint(0, h)
                col = np.random.randint(0, w - 1)
                if chs == 1:
                    values1.append(img[row, col])
                    values2.append(img[row, col + 1])
                else:
                    values1.append(img[row, col, channel])
                    values2.append(img[row, col + 1, channel])
            elif direction == 'vertical':
                row = np.random.randint(0, h - 1)
                col = np.random.randint(0, w)
                if chs == 1:
                    values1.append(img[row, col])
                    values2.append(img[row + 1, col])
                else:
                    values1.append(img[row, col, channel])
                    values2.append(img[row + 1, col, channel])
            elif direction == 'diagonal':
                row = np.random.randint(0, h - 1)
                col = np.random.randint(0, w - 1)
                if chs == 1:
                    values1.append(img[row, col])
                    values2.append(img[row + 1, col + 1])
                else:
                    values1.append(img[row, col, channel])
                    values2.append(img[row + 1, col + 1, channel])
        values1 = np.array(values1)
        values2 = np.array(values2)
        # Add check to avoid division by zero
        stddev1 = np.std(values1)
        stddev2 = np.std(values2)
        if stddev1 == 0 or stddev2 == 0:
            corr = np.nan  # Handle the case where stddev is zero (divide by zero)
        else:
            corr = np.corrcoef(values1, values2)[0, 1]
        result.append(corr)
    return result

def plot_pixel_correlation(img_orig, img_enc, direction='horizontal'):
    color_names = ['Blue', 'Green', 'Red']
    colors = ['b', 'g', 'r']
    plt.figure(figsize=(16, 10))
    for i in range(3):
        x_orig, y_orig = scatter_pixel_pairs(img_orig, i, direction)
        x_enc, y_enc = scatter_pixel_pairs(img_enc, i, direction)
        plt.subplot(3, 2, i * 2 + 1)
        plt.scatter(x_orig, y_orig, color=colors[i], s=1, alpha=0.5)
        plt.subplot(3, 2, i * 2 + 2)
        plt.scatter(x_enc, y_enc, color=colors[i], s=1, alpha=0.5)
    plt.tight_layout()
    plt.show()

def scatter_pixel_pairs(img, channel, direction='horizontal', n_points=3000):
    if len(img.shape) == 2:
        h, w = img.shape
        chs = 1
    else:
        h, w, chs = img.shape
    np.random.seed(10)
    xs = []
    ys = []
    for _ in range(n_points):
        if direction == 'horizontal':
            row = np.random.randint(0, h)
            col = np.random.randint(0, w - 1)
            if chs == 1:
                x = img[row, col]
                y = img[row, col + 1]
            else:
                x = img[row, col, channel]
                y = img[row, col + 1, channel]
        elif direction == 'vertical':
            row = np.random.randint(0, h - 1)
            col = np.random.randint(0, w)
            if chs == 1:
                x = img[row, col]
                y = img[row + 1, col]
            else:
                x = img[row, col, channel]
                y = img[row + 1, col, channel]
        elif direction == 'diagonal':
            row = np.random.randint(0, h - 1)
            col = np.random.randint(0, w - 1)
            if chs == 1:
                x = img[row, col]
                y = img[row + 1, col + 1]
            else:
                x = img[row, col, channel]
                y = img[row + 1, col + 1, channel]
        xs.append(x)
        ys.append(y)
    return np.array(xs), np.array(ys)

def plot_histograms(original_img, encrypted_img):
    if len(original_img.shape) == 2:
        # Grayscale or binary
        plt.figure(figsize=(12, 5))
        plt.subplot(1, 2, 1)
        plt.hist(original_img.flatten(), bins=256, color='black', alpha=0.7)
        plt.subplot(1, 2, 2)
        plt.hist(encrypted_img.flatten(), bins=256, color='black', alpha=0.7)
        plt.tight_layout()
        plt.show()
    else:
        color = ['b', 'g', 'r']
        plt.figure(figsize=(16, 8))
        for i, col in enumerate(color):
            plt.subplot(2, 3, i+1)
            plt.hist(original_img[:, :, i].flatten(), bins=256, color=col, alpha=0.7)
            plt.subplot(2, 3, i+4)
            plt.hist(encrypted_img[:, :, i].flatten(), bins=256, color=col, alpha=0.7)
        plt.tight_layout()
        plt.show()

# ---------- Main ----------
def main():
    img_path = 'white.png'  # ضع مسار الصورة هنا
    encrypted_img, original_img = encrypt_image_with_dna(img_path)

    # الحسابات الأساسية
    entropy_original = calculate_entropy(original_img)
    entropy_encrypted = calculate_entropy(encrypted_img)
    psnr_val = PSNR(original_img, encrypted_img)
    npcr_val = NPCR(original_img, encrypted_img)
    uaci_val = UACI(original_img, encrypted_img)
    directions = ['horizontal', 'vertical', 'diagonal']
    corr_orig = {dir: correlation_coeff(original_img, dir) for dir in directions}
    corr_enc = {dir: correlation_coeff(encrypted_img, dir) for dir in directions}

    # بناء جدول النتائج
    data = [
        ["Entropy", f"{entropy_original:.4f}", f"{entropy_encrypted:.4f}"],
        ["PSNR", f"{psnr_val:.4f}", "-"],
        ["NPCR (%)", f"{npcr_val:.4f}", "-"],
        ["UACI (%)", f"{uaci_val:.4f}", "-"],
    ]
    for dir in directions:
        data.append([f"Correlation {dir.capitalize()} (B)", f"{corr_orig[dir][0]:.4f}", f"{corr_enc[dir][0]:.4f}"])
        data.append([f"Correlation {dir.capitalize()} (G)", f"{corr_orig[dir][1]:.4f}", f"{corr_enc[dir][1]:.4f}"])
        data.append([f"Correlation {dir.capitalize()} (R)", f"{corr_orig[dir][2]:.4f}", f"{corr_enc[dir][2]:.4f}"])

    df = pd.DataFrame(data, columns=["Metric", "Original", "Encrypted"])
    print('\n=== Image Encryption Analysis Results ===\n')
    print(df.to_string(index=False))

    # حفظ النتائج في ملف وورد
    save_results_to_word(df, "results.docx")

    # رسم Histogram لجميع القنوات
    plot_histograms(original_img, encrypted_img)

    # رسم Correlation scatter لكل لون في كل اتجاه (الأصلية والمشفرة)
    for direction in directions:
        plot_pixel_correlation(original_img, encrypted_img, direction)

if __name__ == "__main__":
    main()
